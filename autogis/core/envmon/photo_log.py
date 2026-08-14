"""photo_log.py — photographic log appendix from harvest photo metadata.

The standard consulting deliverable (photo #, thumbnail, date, direction,
coordinates, blank description column for hand-editing) in three formats:
xlsx (openpyxl, mirrors Tool 7.4's embedding path), html
(``report_html``), docx (python-docx via the ``report-docx`` extra).
Image and docx libraries are lazy-imported. No arcpy. No arcgis.
"""
from __future__ import annotations

import base64
import io
from pathlib import Path

from autogis.core.envmon.photo_metadata import PhotoRecord
from autogis.core.envmon.well_inspection_photo_report import (
    prepare_image_bytes)

DOCX_HINT = ("python-docx is required for --format docx; install with: "
             "pip install \"autogis[report-docx]\"")
_CARDINALS = ["N", "NE", "E", "SE", "S", "SW", "W", "NW"]
_THUMB_BOX = (800, 800)
_XLSX_BOX = (300, 225)


def _cardinal(deg: float) -> str:
    return _CARDINALS[round(deg / 45.0) % 8]


def _direction(r: PhotoRecord) -> str:
    if r.heading_deg is None:
        return ""
    ref = " (magnetic)" if r.heading_ref == "M" else ""
    return f"{r.heading_deg:.0f}° {_cardinal(r.heading_deg)}{ref}"


def _coords(r: PhotoRecord) -> str:
    if r.exif_lat is None or r.exif_lon is None:
        return ""
    return f"{r.exif_lat:.6f}, {r.exif_lon:.6f}"


def _feature(r: PhotoRecord) -> str:
    oid = f" OID {r.objectid}" if r.objectid is not None else ""
    return f"{r.group}{oid}"


def _xlsx_text(value: str) -> str:
    """Keep untrusted text from becoming an Excel formula."""
    return "'" + value if value.startswith(("=", "+", "-", "@")) else value


def _bytes_of(r: PhotoRecord, box) -> bytes | None:
    p = Path(r.saved_path)
    return prepare_image_bytes(p, box) if p.is_file() else None


def write_log(records: list[PhotoRecord], out_path: Path, *,
              fmt: str = "xlsx", title: str = "Photographic Log") -> int:
    writer = {"xlsx": _write_xlsx, "html": _write_html,
              "docx": _write_docx}.get(fmt)
    if writer is None:
        raise ValueError(f"unknown log format: {fmt!r} (xlsx|html|docx)")
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    writer(records, out_path, title)
    return len(records)


def _write_xlsx(records, out_path, title):
    from openpyxl import Workbook
    from openpyxl.drawing.image import Image as XLImage
    from openpyxl.styles import Alignment, Font
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "Photo Log"
    headers = ["Photo #", "Image", "Group / Feature", "Taken", "Direction",
               "Coordinates", "Description", "Source Path"]
    ws.append(headers)
    for c, _ in enumerate(headers, start=1):
        ws.cell(row=1, column=c).font = Font(bold=True)
    widths = [8, 44, 28, 20, 16, 24, 40, 50]
    for c, wd in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(c)].width = wd
    wrap = Alignment(wrap_text=True, vertical="top")
    for i, r in enumerate(records, start=1):
        row_no = i + 1
        ws.cell(row=row_no, column=1, value=i)
        data = _bytes_of(r, _XLSX_BOX)
        if data is not None:
            img = XLImage(io.BytesIO(data))
            ws.add_image(img, f"B{row_no}")
            ws.row_dimensions[row_no].height = _XLSX_BOX[1] * 0.75
        for col, val in ((3, _feature(r)), (4, r.taken_at or ""),
                         (5, _direction(r)), (6, _coords(r)), (7, ""),
                         (8, r.saved_path)):
            ws.cell(row=row_no, column=col,
                    value=_xlsx_text(val)).alignment = wrap
    wb.save(out_path)


def _write_html(records, out_path, title):
    from autogis.core.common import report_html as rh

    images, rows = [], []
    for i, r in enumerate(records, start=1):
        data = _bytes_of(r, _THUMB_BOX)
        caption = (f"Photo {i} — {_feature(r)}"
                   + (f" — {r.taken_at}" if r.taken_at else "")
                   + (f" — {_direction(r)}" if r.heading_deg is not None
                      else ""))
        if data is not None:
            src = ("data:image/jpeg;base64,"
                   + base64.b64encode(data).decode("ascii"))
            images.append((src, caption))
        rows.append([i, _feature(r), r.taken_at or "", _direction(r),
                     _coords(r), "", Path(r.saved_path).name])
    out_path.write_text(
        rh.render_document(title=title, sections=[
            rh.section("Photos", rh.photo_grid(images)),
            rh.section("Index", rh.table(
                ["Photo #", "Group / Feature", "Taken", "Direction",
                 "Coordinates", "Description", "File"], rows)),
        ]), encoding="utf-8")


def _write_docx(records, out_path, title):
    try:
        import docx
        from docx.shared import Inches
    except ImportError as exc:
        raise ImportError(DOCX_HINT) from exc
    doc = docx.Document()
    doc.add_heading(title, level=1)
    for i, r in enumerate(records, start=1):
        data = _bytes_of(r, _THUMB_BOX)
        if data is not None:
            doc.add_picture(io.BytesIO(data), width=Inches(4.5))
        meta = [f"Photo {i} — {_feature(r)}"]
        if r.taken_at:
            meta.append(f"Taken: {r.taken_at}")
        if r.heading_deg is not None:
            meta.append(f"Direction: {_direction(r)}")
        if _coords(r):
            meta.append(f"Coordinates: {_coords(r)}")
        p = doc.add_paragraph("\n".join(meta))
        p.runs[0].bold = True
        doc.add_paragraph("Description: ")
        doc.add_paragraph()
    doc.save(str(out_path))
